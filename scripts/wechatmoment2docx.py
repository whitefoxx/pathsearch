"""背景：一个学长希望我帮他将朋友圈内容导入到word文档中
首先，我已经通过微信书导出了他的数据，剩下的就是使用python-docx工具将数据写入word文档
"""

from ast import expr_context
import json
import hashlib
import os
from pydoc import doc
from PIL import Image
from numpy import add

import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError


def download_moments():
    base_url = "https://weixinshu.com/api/partner/wxbook/items?weixin_id=wxid_8300043079512&page=%d&page_size=100"
    for i in range(1, 61):
        url = base_url % i
        r = requests.get(url)
        with open("./moments/page%d.json" % i, "w") as f:
            f.write(json.dumps(r.json(), ensure_ascii=False))


def download_images():
    for i in range(1, 61):
        with open("./moments/page%d.json" % i, "r") as f:
            data = json.load(f)["data"]
            for msg in data.get("messages", []):
                if "img_urls" in msg:
                    for url in msg["img_urls"]:
                        print(url)
                        name = hashlib.md5(url.encode()).hexdigest() + ".jpg"
                        r = requests.get(url)
                        with open("./images/%s" % name, "wb") as f:
                            f.write(r.content)


def get_months():
    months = set()
    for i in range(1, 61):
        with open("./moments/page%d.json" % i, "r") as f:
            data = json.load(f)["data"]
            for msg in data.get("messages", []):
                post_date = msg["post_date"]
                month = post_date.rsplit("-", 1)[0]
                months.add(month)
    months = list(months)
    months.sort()
    for month in months:
        print(month)
    return months


def get_month_data(month):
    messages = []
    for i in range(1, 61):
        with open("./moments/page%d.json" % i, "r") as f:
            data = json.load(f)["data"]
            for msg in data.get("messages", []):
                post_date = msg["post_date"]
                m = post_date.rsplit("-", 1)[0]
                if m == month:
                    messages.append(msg)
    messages.reverse()
    return messages


def image_to_jpg(url, fpath):
    try:
        img = Image.open(fpath)
    except Image.UnidentifiedImageError:
        print("download %s" % url)
        r = requests.get(url)
        with open(fpath, "wb") as f:
            f.write(r.content)
        img = Image.open(fpath)
    img.convert("RGB").save(fpath)


def write_month(messages, month):
    def add_picture(document, fname):
        if os.path.exists(fname):
            document.add_picture(fname, width=Inches(6))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph()

    if os.path.exists("./docs/%s.docx" % month):
        return
    document = Document()
    document.add_heading(month, 0)
    for msg in messages:
        post_date = msg["post_date"]
        document.add_heading(post_date, 1)
        content = msg.get("content")
        if content:
            paragraph = document.add_paragraph(content)
            paragraph.aligment = WD_ALIGN_PARAGRAPH.LEFT
        for url in msg.get("img_urls", []):
            try:
                fname = "./images/" + hashlib.md5(url.encode()).hexdigest() + ".jpg"
                add_picture(document, fname)
            except UnrecognizedImageError as e:
                try:
                    image_to_jpg(url, fname)
                    add_picture(document, fname)
                except Exception as e:
                    print(url)
                    print(fname)
                    print("Unkonw error1: %s" % e)
            except Exception as e:
                print(url)
                print(fname)
                print("Unkonw error2: %s" % e)
    document.save("./docs/%s.docx" % month)


# download_moments()
# download_images()
months = get_months()
for month in months:
    print(month)
    messages = get_month_data(month)
    write_month(messages, month)
# month = "2015-04"
# messages = get_month_data(month)
# write_month(messages, month)
